"""
DOCX to PDF Converter — python-docx → HTML → Playwright PDF

Key behaviour:
  • Reads page size & margins directly from the DOCX section so the
    available content area is exact.
  • After rendering, measures actual content height in the browser.
    If the content overflows (would spill onto an extra page) it applies
    a CSS zoom-out so everything fits in the correct number of pages.
  • Strips any blank trailing pages Chromium may append.
"""

import os
import logging
import tempfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from playwright.sync_api import sync_playwright
from pypdf import PdfReader, PdfWriter

SUPPORTED_FORMATS = {'.docx', '.doc'}
logger = logging.getLogger(__name__)

# ── Tiny helpers ──────────────────────────────────────────────────────────────

def _esc(t):
    return t.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;').replace('"','&quot;')

def _has_bottom_border(para):
    pPr = para._p.find(qn('w:pPr'))
    if pPr is None: return False
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None: return False
    return pBdr.find(qn('w:bottom')) is not None

def _align_css(para):
    a = para.alignment
    if a == WD_ALIGN_PARAGRAPH.CENTER:  return 'center'
    if a == WD_ALIGN_PARAGRAPH.RIGHT:   return 'right'
    if a == WD_ALIGN_PARAGRAPH.JUSTIFY: return 'justify'
    return 'left'

def _hyperlink_map(doc):
    m = {}
    try:
        for rel in doc.part.rels.values():
            if 'hyperlink' in rel.reltype:
                m[rel.rId] = rel._target
    except Exception:
        pass
    return m

# ── Run → HTML ────────────────────────────────────────────────────────────────

def _render_run(run):
    text = _esc(run.text)
    if not text:
        return ''
    styles = []
    try:
        if run.font.size:
            styles.append(f"font-size:{run.font.size.pt:.1f}pt")
    except Exception: pass
    try:
        if run.bold:      styles.append('font-weight:bold')
    except Exception: pass
    try:
        if run.italic:    styles.append('font-style:italic')
    except Exception: pass
    try:
        if run.underline: styles.append('text-decoration:underline')
    except Exception: pass
    try:
        c = run.font.color
        if c and c.rgb:
            hex_c = f"#{c.rgb}"
            if hex_c.lower() not in ('#000000','#auto'):
                styles.append(f"color:{hex_c}")
    except Exception: pass
    try:
        fn = run.font.name
        if fn: styles.append(f"font-family:'{fn}',Calibri,Arial,sans-serif")
    except Exception: pass

    s = ';'.join(styles)
    return f'<span style="{s}">{text}</span>' if s else text

# ── Paragraph → HTML ──────────────────────────────────────────────────────────

def _render_paragraph(para, hmap):
    # detect list
    numPr  = para._p.find(f'.//{qn("w:numPr")}')
    is_list = numPr is not None

    # inner content (runs + hyperlinks)
    inner = ''
    for child in para._p:
        if child.tag == qn('w:hyperlink'):
            rid  = child.get(qn('r:id'))
            url  = hmap.get(rid, '')
            li   = ''
            for r_el in child.findall(qn('w:r')):
                from docx.text.run import Run
                li += _render_run(Run(r_el, para))
            inner += (f'<a href="{_esc(url)}" style="color:#1155CC;'
                      f'text-decoration:underline">{li}</a>' if url else li)
        elif child.tag == qn('w:r'):
            from docx.text.run import Run
            inner += _render_run(Run(child, para))

    if not inner:
        inner = '&nbsp;'

    # paragraph styles
    ps = []
    al = _align_css(para)
    if al != 'left':
        ps.append(f'text-align:{al}')

    # spacing — cap at 12pt to avoid runaway gaps
    try:
        sb = para.paragraph_format.space_before
        if sb and sb.pt > 0:
            ps.append(f'margin-top:{min(sb.pt, 12):.1f}pt')
    except Exception: pass
    try:
        sa = para.paragraph_format.space_after
        if sa and sa.pt > 0:
            ps.append(f'margin-bottom:{min(sa.pt, 8):.1f}pt')
    except Exception: pass

    # indent (non-list only)
    if not is_list:
        try:
            li_val = para.paragraph_format.left_indent
            if li_val and li_val.pt > 0:
                ps.append(f'padding-left:{li_val.pt:.1f}pt')
        except Exception: pass

    style = f' style="{";".join(ps)}"' if ps else ''
    border = ('<hr style="border:none;border-bottom:1px solid #000;'
              'margin:1pt 0 3pt 0"/>') if _has_bottom_border(para) else ''

    if is_list:
        return f'<li{style}>{inner}</li>', True
    return f'<p{style}>{inner}</p>{border}', False

# ── Table → HTML ──────────────────────────────────────────────────────────────

def _render_table(table, hmap):
    rows = ''
    for i, row in enumerate(table.rows):
        cells = ''
        for cell in row.cells:
            body = ''.join(_render_paragraph(p, hmap)[0] for p in cell.paragraphs)
            tag = 'th' if i == 0 else 'td'
            cells += f'<{tag}>{body}</{tag}>'
        rows += f'<tr>{cells}</tr>'
    return f'<table>{rows}</table>'

# ── DOCX → HTML ───────────────────────────────────────────────────────────────

def _docx_to_html(input_path: Path, margins: dict) -> str:
    doc   = Document(str(input_path))
    hmap  = _hyperlink_map(doc)
    body  = ''
    in_ul = False

    for block in doc.element.body:
        tag = block.tag
        if tag == qn('w:p'):
            from docx.text.paragraph import Paragraph
            para    = Paragraph(block, doc)
            html, is_li = _render_paragraph(para, hmap)
            if is_li and not in_ul:
                body += '<ul>'; in_ul = True
            elif not is_li and in_ul:
                body += '</ul>'; in_ul = False
            body += html
        elif tag == qn('w:tbl'):
            if in_ul:
                body += '</ul>'; in_ul = False
            from docx.table import Table
            body += _render_table(Table(block, doc), hmap)

    if in_ul:
        body += '</ul>'

    # default font from Normal style
    dfont = 'Calibri,"Segoe UI",Arial,sans-serif'
    dsize = '11pt'
    try:
        n = doc.styles['Normal']
        if n.font.name:  dfont = f"'{n.font.name}',Calibri,Arial,sans-serif"
        if n.font.size:  dsize = f"{n.font.size.pt:.1f}pt"
    except Exception: pass

    mt = margins['top'];   mb = margins['bottom']
    ml = margins['left'];  mr = margins['right']

    return f"""<!DOCTYPE html>
<html>
<head><meta charset="utf-8"/>
<style>
  @page {{
    size: A4;
    margin: {mt:.1f}pt {mr:.1f}pt {mb:.1f}pt {ml:.1f}pt;
  }}
  *,*::before,*::after{{box-sizing:border-box;margin:0;padding:0}}
  html,body{{
    background:white;
    font-family:{dfont};
    font-size:{dsize};
    line-height:1.15;
    color:#000;
  }}
  /* remove trailing margin — main guard against spurious blank page */
  body>*:last-child,body>*:last-child *:last-child{{
    margin-bottom:0!important;padding-bottom:0!important;
  }}
  p{{margin:0;padding:0}}
  ul{{margin:0;padding-left:22pt;list-style-type:disc}}
  li{{margin-bottom:1pt}}
  a{{color:#1155CC;text-decoration:underline}}
  table{{border-collapse:collapse;width:100%;margin:4pt 0;font-size:10pt}}
  th,td{{border:1px solid #ccc;padding:3pt 5pt;text-align:left;vertical-align:top}}
  th{{background:#2c3e50;color:white;font-weight:bold}}
  tr:nth-child(even){{background:#f5f7fa}}
  hr{{border:none;border-bottom:1px solid #000;margin:1pt 0 3pt 0}}
</style>
</head>
<body>{body}</body>
</html>"""

# ── Page size from DOCX section ───────────────────────────────────────────────

def _read_margins(doc) -> dict:
    """Read page margins (in pt) from the first section."""
    defaults = dict(top=72, bottom=72, left=72, right=72,
                    page_h=841.89, page_w=595.28)
    try:
        s = doc.sections[0]
        return dict(
            top    = s.top_margin.pt,
            bottom = s.bottom_margin.pt,
            left   = s.left_margin.pt,
            right  = s.right_margin.pt,
            page_h = s.page_height.pt,
            page_w = s.page_width.pt,
        )
    except Exception:
        return defaults

# ── Playwright export with auto-scale ────────────────────────────────────────

def _html_to_pdf(html: str, output_path: str, tmp_dir: Path, margins: dict):
    html_path = tmp_dir / 'document.html'
    html_path.write_text(html, encoding='utf-8')

    # Available content height in CSS px (96 dpi: 1pt = 96/72 px)
    PT_TO_PX     = 96 / 72
    available_pt = margins['page_h'] - margins['top'] - margins['bottom']
    available_px = available_pt * PT_TO_PX

    with sync_playwright() as pw:
        browser = pw.chromium.launch(args=['--no-sandbox', '--disable-gpu'])
        page    = browser.new_page()
        page.goto(f'file://{html_path}', wait_until='networkidle')
        page.wait_for_timeout(400)

        # Measure actual rendered height
        content_px = page.evaluate('document.body.scrollHeight')
        logger.info(f"Content height: {content_px:.0f}px  |  Available: {available_px:.0f}px")

        if content_px > available_px:
            scale = available_px / content_px
            # Zoom out — Chromium supports the CSS zoom property
            page.evaluate(f'document.body.style.zoom = "{scale:.6f}"')
            logger.info(f"Auto-scaled content to {scale*100:.1f}% to fit page")

        page.pdf(path=output_path, format='A4', print_background=True)
        browser.close()

    logger.info(f"PDF written: {output_path}")

# ── Strip blank trailing pages ────────────────────────────────────────────────

def _strip_blank_trailing_pages(pdf_path: str):
    reader = PdfReader(pdf_path)
    pages  = list(reader.pages)
    before = len(pages)

    while pages:
        last     = pages[-1]
        text     = (last.extract_text() or '').strip()
        raw_size = 0
        if '/Contents' in last:
            c = last['/Contents']
            if hasattr(c, 'get_data'):
                raw_size = len(c.get_data())
            elif isinstance(c, list):
                raw_size = sum(len(x.get_data()) for x in c if hasattr(x, 'get_data'))
        if text == '' and raw_size < 100:
            pages.pop()
            logger.info("Stripped blank trailing page.")
        else:
            break

    if len(pages) == before:
        return

    writer = PdfWriter()
    for p in pages:
        writer.add_page(p)
    with open(pdf_path, 'wb') as f:
        writer.write(f)

# ── Main entry point ──────────────────────────────────────────────────────────

def convert_docx_to_pdf(input_path: str, output_path: str) -> dict:
    ext = Path(input_path).suffix.lower()
    if ext not in SUPPORTED_FORMATS:
        return {'success': False,
                'message': f"Unsupported format '{ext}'. Supported: {', '.join(SUPPORTED_FORMATS)}"}

    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp = Path(tmp_dir)
        try:
            logger.info("Step 1: Reading DOCX (python-docx)")
            doc     = Document(str(input_path))
            margins = _read_margins(doc)
            logger.info(f"  Margins (pt): top={margins['top']:.0f} bottom={margins['bottom']:.0f} "
                        f"left={margins['left']:.0f} right={margins['right']:.0f}")

            logger.info("Step 2: DOCX → HTML")
            html = _docx_to_html(Path(input_path), margins)

            logger.info("Step 3: HTML → PDF (Playwright + auto-scale)")
            _html_to_pdf(html, output_path, tmp, margins)

            logger.info("Step 4: Strip blank trailing pages")
            _strip_blank_trailing_pages(output_path)

        except Exception as e:
            logger.exception("Conversion failed")
            return {'success': False, 'message': f'Conversion failed: {e}'}

    size_kb = os.path.getsize(output_path) / 1024
    pages   = len(PdfReader(output_path).pages)
    logger.info(f"Output: {pages} page(s), {size_kb:.1f} KB")
    return {'success': True,
            'message': f'Converted {pages} page(s). Output: {size_kb:.1f} KB'}
