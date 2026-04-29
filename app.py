"""
PDF Converter - Flask Application
Converts DOCX, Images, and Excel files to PDF.
Supports single and multiple file uploads.
"""

import os
import uuid
import logging
from pathlib import Path
from datetime import datetime
from dotenv import load_dotenv


from pypdf import PdfWriter, PdfReader

from flask import (
    Flask, request, jsonify, send_file,
    render_template, abort
)
from werkzeug.utils import secure_filename

from converters.image_converter import convert_image_to_pdf, SUPPORTED_FORMATS as IMAGE_FORMATS
from converters.doc_converter import convert_docx_to_pdf, SUPPORTED_FORMATS as DOC_FORMATS
from converters.excel_converter import convert_excel_to_pdf, SUPPORTED_FORMATS as EXCEL_FORMATS
from font_manager import get_font_path

# ── Configuration ─────────────────────────────────────────────────────────────
load_dotenv()
BASE_DIR = Path(__file__).parent
UPLOAD_FOLDER = BASE_DIR / 'uploads'
OUTPUT_FOLDER = BASE_DIR / 'outputs'
MAX_CONTENT_LENGTH = 200 * 1024 * 1024  # 200 MB total (multi-file)

UPLOAD_FOLDER.mkdir(exist_ok=True)
OUTPUT_FOLDER.mkdir(exist_ok=True)

ALL_ALLOWED = IMAGE_FORMATS | DOC_FORMATS | EXCEL_FORMATS

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s  %(levelname)-8s  %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S',
)
logger = logging.getLogger(__name__)

# ── App ───────────────────────────────────────────────────────────────────────

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH
app.secret_key = os.environ.get('SECRET_KEY', 'dev-secret-change-in-production')


# ── Helpers ───────────────────────────────────────────────────────────────────

def allowed_file(filename: str) -> bool:
    return Path(filename).suffix.lower() in ALL_ALLOWED


def get_converter(ext: str):
    if ext in IMAGE_FORMATS:
        return convert_image_to_pdf
    if ext in DOC_FORMATS:
        return convert_docx_to_pdf
    if ext in EXCEL_FORMATS:
        return convert_excel_to_pdf
    return None


def cleanup_old_files(folder: Path, max_age_seconds: int = 3600):
    now = datetime.now().timestamp()
    removed = 0
    for f in folder.iterdir():
        if f.is_file() and (now - f.stat().st_mtime) > max_age_seconds:
            try:
                f.unlink()
                removed += 1
            except OSError:
                pass
    if removed:
        logger.info(f"Cleaned up {removed} old file(s) from {folder}")


# ── Routes ────────────────────────────────────────────────────────────────────

@app.route('/')
def index():
    return render_template('index.html',
                           image_formats=sorted(IMAGE_FORMATS),
                           doc_formats=sorted(DOC_FORMATS),
                           excel_formats=sorted(EXCEL_FORMATS))


@app.route('/convert', methods=['POST'])
def convert():
    """
    POST /convert
    Form field : files[]  (one or more multipart files)

    Single file  -> { success, message, download_id, filename }
    Multi  files -> { success, message, download_id, filename, results[], is_zip }
    """
    cleanup_old_files(UPLOAD_FOLDER)
    cleanup_old_files(OUTPUT_FOLDER)

    files = request.files.getlist('files[]')
    files = [f for f in files if f.filename]

    if not files:
        return jsonify({'success': False, 'error': 'No files selected'}), 400

    # Validate all extensions upfront
    for f in files:
        if not allowed_file(f.filename):
            ext = Path(f.filename).suffix.lower()
            return jsonify({
                'success': False,
                'error': f"'{f.filename}': type '{ext}' is not supported. "
                         f"Allowed: {', '.join(sorted(ALL_ALLOWED))}"
            }), 415

    uid = str(uuid.uuid4())
    results = []
    converted_paths = []  # (Path, pdf_filename)
    used_pdf_names = set()  # track within this batch

    for file in files:
        safe_name = secure_filename(file.filename)
        input_path = UPLOAD_FOLDER / f"{uid}_{safe_name}"
        file.save(input_path)

        ext = Path(safe_name).suffix.lower()
        base_pdf = Path(safe_name).stem + '.pdf'
        # Deduplicate within this batch (e.g. report.png + report.docx → report.pdf clash)
        if base_pdf in used_pdf_names:
            base_pdf = Path(safe_name).stem + f'_{ext.lstrip(".")}.pdf'
        used_pdf_names.add(base_pdf)
        pdf_name = base_pdf
        output_path = OUTPUT_FOLDER / f"{uid}_{pdf_name}"

        logger.info(f"Converting [{ext}] -> PDF  |  {safe_name}")
        converter = get_converter(ext)
        result = converter(str(input_path), str(output_path))

        try:
            input_path.unlink()
        except OSError:
            pass

        entry = {
            'filename': safe_name,
            'output': pdf_name,
            'success': result['success'],
            'message': result['message'],
            'download_id': f"{uid}_{pdf_name}" if result['success'] else None,
        }
        results.append(entry)

        if result['success']:
            converted_paths.append((output_path, pdf_name))

    # ── Merge all successful PDFs into one ────────────────────────────────────
    successful = [r for r in results if r['success']]
    failed     = [r for r in results if not r['success']]

    if not successful:
        return jsonify({
            'success': False,
            'error': 'All conversions failed.',
            'results': results,
        }), 500

    merged_name = f"merged_{uid}.pdf"
    merged_path = OUTPUT_FOLDER / merged_name

    from pypdf import PdfWriter, PdfReader
    writer = PdfWriter()
    for pdf_path, _ in converted_paths:
        reader = PdfReader(str(pdf_path))
        for page in reader.pages:
            writer.add_page(page)
        # clean up individual intermediate PDF
        try:
            pdf_path.unlink()
        except OSError:
            pass

    with open(merged_path, 'wb') as f:
        writer.write(f)

    size_kb = merged_path.stat().st_size / 1024
    summary = f"{len(successful)} file(s) merged into one PDF ({size_kb:.1f} KB)."
    if failed:
        summary += f" Skipped (failed): {', '.join(r['filename'] for r in failed)}"

    logger.info(summary)
    return jsonify({
        'success': True,
        'message': summary,
        'download_id': merged_name,
        'filename': 'merged.pdf',
        'results': results,
        'pages': sum(len(PdfReader(str(p)).pages) for p, _ in []) ,  # already merged
    })


@app.route('/translate', methods=['POST'])
def translate_file():
    """
    POST /translate
    Form fields: file (PDF / DOCX / DOC), language (target language name)
    Returns: { success, message, download_id, filename }
    """
    cleanup_old_files(UPLOAD_FOLDER)
    cleanup_old_files(OUTPUT_FOLDER)

    if 'file' not in request.files:
        return jsonify({'success': False, 'error': 'No file provided'}), 400

    file = request.files['file']
    if not file.filename:
        return jsonify({'success': False, 'error': 'No file selected'}), 400

    ext = Path(file.filename).suffix.lower()
    SUPPORTED_TRANSLATE = {'.pdf', '.docx', '.doc'}
    if ext not in SUPPORTED_TRANSLATE:
        return jsonify({
            'success': False,
            'error': f"'{ext}' is not supported for translation. Supported: PDF, DOCX, DOC"
        }), 415

    target_language = request.form.get('language', 'Spanish').strip()

    uid = str(uuid.uuid4())
    safe_name = secure_filename(file.filename)
    input_path = UPLOAD_FOLDER / f"{uid}_{safe_name}"
    file.save(input_path)

    try:
        # ── Google Translate setup (via deep-translator) ─────────────────
        from deep_translator import GoogleTranslator

        LANG_CODES = {
            'afrikaans': 'af', 'albanian': 'sq', 'amharic': 'am',
            'arabic': 'ar', 'armenian': 'hy', 'azerbaijani': 'az',
            'basque': 'eu', 'belarusian': 'be', 'bengali': 'bn',
            'bosnian': 'bs', 'bulgarian': 'bg', 'catalan': 'ca',
            'chinese': 'zh-cn', 'chinese simplified': 'zh-cn',
            'chinese traditional': 'zh-tw', 'croatian': 'hr',
            'czech': 'cs', 'danish': 'da', 'dutch': 'nl',
            'english': 'en', 'esperanto': 'eo', 'estonian': 'et',
            'finnish': 'fi', 'french': 'fr', 'galician': 'gl',
            'georgian': 'ka', 'german': 'de', 'greek': 'el',
            'gujarati': 'gu', 'haitian creole': 'ht', 'hausa': 'ha',
            'hebrew': 'iw', 'hindi': 'hi', 'hungarian': 'hu',
            'icelandic': 'is', 'igbo': 'ig', 'indonesian': 'id',
            'irish': 'ga', 'italian': 'it', 'japanese': 'ja',
            'javanese': 'jw', 'kannada': 'kn', 'kazakh': 'kk',
            'khmer': 'km', 'korean': 'ko', 'kurdish': 'ku',
            'kyrgyz': 'ky', 'lao': 'lo', 'latin': 'la',
            'latvian': 'lv', 'lithuanian': 'lt', 'luxembourgish': 'lb',
            'macedonian': 'mk', 'malagasy': 'mg', 'malay': 'ms',
            'malayalam': 'ml', 'maltese': 'mt', 'maori': 'mi',
            'marathi': 'mr', 'mongolian': 'mn', 'myanmar': 'my',
            'nepali': 'ne', 'norwegian': 'no', 'odia': 'or',
            'pashto': 'ps', 'persian': 'fa', 'polish': 'pl',
            'portuguese': 'pt', 'punjabi': 'pa', 'romanian': 'ro',
            'russian': 'ru', 'samoan': 'sm', 'scots gaelic': 'gd',
            'serbian': 'sr', 'sesotho': 'st', 'shona': 'sn',
            'sindhi': 'sd', 'sinhala': 'si', 'slovak': 'sk',
            'slovenian': 'sl', 'somali': 'so', 'spanish': 'es',
            'sundanese': 'su', 'swahili': 'sw', 'swedish': 'sv',
            'tajik': 'tg', 'tamil': 'ta', 'telugu': 'te',
            'thai': 'th', 'turkish': 'tr', 'ukrainian': 'uk',
            'urdu': 'ur', 'uyghur': 'ug', 'uzbek': 'uz',
            'vietnamese': 'vi', 'welsh': 'cy', 'xhosa': 'xh',
            'yiddish': 'yi', 'yoruba': 'yo', 'zulu': 'zu',
        }

        lang_code = LANG_CODES.get(target_language.lower(), target_language.lower())
        translator = GoogleTranslator(source='auto', target=lang_code)
        MAX_CHUNK = 4500

        def gt(text: str) -> str:
            """Translate text via Google Translate (deep-translator)."""
            if not text or not text.strip():
                return text
            return GoogleTranslator(source='auto', target=lang_code).translate(text)

        def translate_block(text: str) -> str:
            """Chunk-translate a large text block."""
            if not text.strip():
                return text
            lines_in = text.splitlines(keepends=True)
            result_parts = []
            buffer = ''
            for line in lines_in:
                if len(buffer) + len(line) > MAX_CHUNK and buffer:
                    result_parts.append(gt(buffer))
                    buffer = line
                else:
                    buffer += line
            if buffer.strip():
                result_parts.append(gt(buffer))
            return ''.join(result_parts)

        stem = Path(safe_name).stem

        # ════════════════════════════════════════════════════════════════════
        #  DOCX / DOC path
        # ════════════════════════════════════════════════════════════════════
        if ext in ('.docx', '.doc'):
            from docx import Document as DocxDocument
            from copy import deepcopy

            doc = DocxDocument(str(input_path))

            def translate_paragraph(para):
                """Translate a paragraph, preserving the first run's formatting."""
                full_text = para.text
                if not full_text.strip():
                    return
                translated = gt(full_text)
                # Clear all runs and replace with translated text in first run
                if para.runs:
                    first_run = para.runs[0]
                    first_run.text = translated
                    for run in para.runs[1:]:
                        run.text = ''
                else:
                    para.add_run(translated)

            # Translate body paragraphs
            for para in doc.paragraphs:
                translate_paragraph(para)

            # Translate table cells
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            translate_paragraph(para)

            # Translate headers and footers
            for section in doc.sections:
                for para in section.header.paragraphs:
                    translate_paragraph(para)
                for para in section.footer.paragraphs:
                    translate_paragraph(para)

            out_name = f"translated_{uid}_{stem}_{target_language.replace(' ', '_')}.docx"
            output_path = OUTPUT_FOLDER / out_name
            doc.save(str(output_path))

            size_kb = output_path.stat().st_size / 1024
            dl_name = f"translated_{stem}_{target_language.replace(' ', '_')}.docx"
            para_count = len([p for p in doc.paragraphs if p.text.strip()])

            logger.info(f"DOCX translation complete: {safe_name} -> {target_language} ({size_kb:.1f} KB)")
            return jsonify({
                'success': True,
                'message': f'Translated to {target_language} ({size_kb:.1f} KB, {para_count} paragraph(s))',
                'download_id': out_name,
                'filename': dl_name,
            })

        # ════════════════════════════════════════════════════════════════════
        #  PDF path  –  layout-preserving translation via PyMuPDF
        # ════════════════════════════════════════════════════════════════════
        import fitz  # pymupdf

        # RTL languages need right-alignment
        RTL_LANGUAGES = {
            'arabic', 'hebrew', 'urdu', 'persian', 'pashto',
            'sindhi', 'uyghur', 'yiddish',
        }
        is_rtl = target_language.lower() in RTL_LANGUAGES
        text_align = 2 if is_rtl else 0  # 0=left, 2=right

        font_path = get_font_path(target_language)

        out_name    = f"translated_{uid}_{stem}_{target_language.replace(' ', '_')}.pdf"
        output_path = OUTPUT_FOLDER / out_name

        doc = fitz.open(str(input_path))

        # Quick check: does the PDF have any extractable text?
        has_text = any(
            page.get_text("text").strip()
            for page in doc
        )
        if not has_text:
            doc.close()
            return jsonify({
                'success': False,
                'error': 'No extractable text found. The PDF may be scanned or image-based.',
            }), 400

        total_pages = len(doc)

        for page in doc:
            # ── 1. Collect all text lines with position & style ───────────
            text_dict = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)

            lines_info = []
            for block in text_dict.get("blocks", []):
                if block.get("type") != 0:   # 0 = text block; skip images
                    continue
                for line in block.get("lines", []):
                    spans = line.get("spans", [])
                    if not spans:
                        continue
                    line_text = "".join(s["text"] for s in spans)
                    if not line_text.strip():
                        continue

                    # Dominant font size (largest span) and colour
                    font_size  = max(s["size"] for s in spans)
                    color_int  = spans[0]["color"]          # sRGB packed int
                    r = ((color_int >> 16) & 0xFF) / 255
                    g = ((color_int >> 8)  & 0xFF) / 255
                    b = ( color_int        & 0xFF) / 255

                    lines_info.append({
                        "bbox":      line["bbox"],          # (x0,y0,x1,y1)
                        "text":      line_text,
                        "font_size": font_size,
                        "color":     (r, g, b),
                    })

            if not lines_info:
                continue

            # ── 2. Translate every line ───────────────────────────────────
            for item in lines_info:
                try:
                    item["translated"] = gt(item["text"])
                except Exception:
                    item["translated"] = item["text"]   # fallback: keep original

            # ── 3. Redact original text (fill with page background colour) ─
            #       Default background is white; expand rect by 1 pt to cover
            #       any anti-aliasing bleed.
            for item in lines_info:
                rect = fitz.Rect(item["bbox"]) + (-1, -1, 1, 1)
                page.add_redact_annot(rect, fill=(1, 1, 1))   # white fill
            # images=0 → don't touch embedded images
            page.apply_redactions(images=0)

            # ── 4. Register font on this page, then insert translated text ──
            FONT_ALIAS = "TFont"
            page.insert_font(fontname=FONT_ALIAS, fontfile=font_path)

            for item in lines_info:
                translated = (item.get("translated") or "").strip()
                if not translated:
                    continue

                rect      = fitz.Rect(item["bbox"])
                font_size = item["font_size"]
                color     = item["color"]

                # Try to fit at the original size; shrink if the translation
                # is longer than the original and overflows the bbox.
                for size in [font_size,
                             font_size * 0.90,
                             font_size * 0.80,
                             font_size * 0.70,
                             font_size * 0.60]:
                    rc = page.insert_textbox(
                        rect,
                        translated,
                        fontsize=size,
                        fontname=FONT_ALIAS,
                        color=color,
                        align=text_align,
                    )
                    if rc >= 0:     # rc ≥ 0 means all text fitted
                        break
                    # rc < 0 → overflow; try a smaller size

        doc.save(str(output_path), garbage=4, deflate=True)
        doc.close()

        size_kb = output_path.stat().st_size / 1024
        dl_name = f"translated_{stem}_{target_language.replace(' ', '_')}.pdf"

        logger.info(f"PDF translation complete (layout-preserved): "
                    f"{safe_name} -> {target_language} ({size_kb:.1f} KB)")
        return jsonify({
            'success': True,
            'message': (f'Translated to {target_language} '
                        f'({size_kb:.1f} KB, {total_pages} page(s)) — '
                        f'layout preserved'),
            'download_id': out_name,
            'filename':    dl_name,
        })

    except RuntimeError as exc:
        return jsonify({'success': False, 'error': str(exc)}), 500
    except Exception as exc:
        logger.exception('Translation error')
        return jsonify({'success': False, 'error': str(exc)}), 500
    finally:
        try:
            input_path.unlink()
        except OSError:
            pass



@app.route('/download/<path:download_id>')
def download(download_id: str):
    safe_id = secure_filename(download_id)
    if safe_id != download_id:
        abort(400)

    file_path = OUTPUT_FOLDER / safe_id
    if not file_path.exists():
        abort(404)

    # Strip uid prefix to restore original filename
    parts = safe_id.split('_')
    original_name = '_'.join(parts[1:]) if len(parts) > 1 else safe_id
    ext = Path(safe_id).suffix.lower()
    mime = ('application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            if ext == '.docx' else 'application/pdf')

    return send_file(
        file_path,
        as_attachment=True,
        download_name=original_name,
        mimetype=mime,
    )


@app.route('/health')
def health():
    return jsonify({'status': 'ok', 'service': 'pdf-converter'})


@app.errorhandler(413)
def too_large(e):
    return jsonify({'success': False, 'error': 'Upload too large (max 200 MB total)'}), 413


@app.errorhandler(404)
def not_found(e):
    return jsonify({'success': False, 'error': 'Resource not found'}), 404


# ── Entry Point ───────────────────────────────────────────────────────────────

if __name__ == '__main__':
    logger.info("Starting PDF Converter  (development mode)")
    logger.info(f"  Supported: {sorted(ALL_ALLOWED)}")
    app.run(debug=True, host='0.0.0.0', port=5000)